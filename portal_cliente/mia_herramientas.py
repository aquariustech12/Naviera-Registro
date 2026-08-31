# portal_cliente/mia_herramientas.py
import os
import re
import json
import shutil
import requests
from datetime import datetime

# OCR
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from PIL import Image
import pdfplumber

# Django models
from naviera_registro.models import Naviera, Buque, RequisitoBuque, PuntoPBIP, DocumentoEntregable, AnalisisMIA

# ============================================================================
# CONFIGURACIÓN
# ============================================================================

# RAG v3 remoto (Epyc/Sonora) — reemplaza el corpus local mia_conocimiento
RAG_V3_URL = "http://100.112.139.108:8020/buscar"
RAG_V3_TOKEN = "6ac319869b1544b85b02d9745bc413566fd21bf9a13b9974573ebf7f8100b291"

# Memoria de análisis de situaciones (MIA 2.0, Epyc/Sonora) — una sola fuente
# de verdad compartida entre COMPASS y MIA 2.0, agosto 2026.
ANALISIS_MEMORIA_URL = "http://100.112.139.108:8010/api/mia/analisis"

print(f"--- MIA HERRAMIENTAS INICIANDO ---")
print(f"OCR: {'✅' if OCR_AVAILABLE else '❌'}")
print(f"DOCX: {'✅' if DOCX_AVAILABLE else '❌'}")

def _inicializar_chroma_DESACTIVADO():
    """Inicializa Chroma y construye índice BM25 en memoria."""
    global vector_db, _bm25_index, _corpus_texts, _corpus_metadatas

    try:
        vector_db = Chroma(
            persist_directory=CHROMA_PATH,
            embedding_function=embeddings,
            collection_name=COLLECTION_NAME
        )
        count = vector_db._collection.count()
        print(f"✅ Base MIA ({COLLECTION_NAME}): {count} documentos")

        # Construir índice BM25
        if BM25_AVAILABLE and count > 0:
            _construir_bm25()

    except Exception as e:
        print(f"❌ ERROR Chroma: {e}")
        vector_db = None


def _construir_bm25():
    """Construye el índice BM25 desde Chroma. Se ejecuta una sola vez."""
    global _bm25_index, _corpus_texts, _corpus_metadatas

    if _bm25_index is not None:
        return

    if vector_db is None:
        return

    try:
        all_docs = vector_db.get()
        _corpus_texts = all_docs['documents']
        _corpus_metadatas = all_docs['metadatas']

        tokenized = [re.findall(r'\b\w+\b', doc.lower()) for doc in _corpus_texts]
        _bm25_index = BM25Okapi(tokenized)

        print(f"✅ Índice BM25 construido: {len(_corpus_texts)} documentos")
    except Exception as e:
        print(f"❌ Error construyendo BM25: {e}")
        _bm25_index = None


# Corpus local (mia_conocimiento) DESACTIVADO — agosto 2026.
# Todas las consultas normativas (PBIP, legislación mexicana, ISM, MARPOL, etc.)
# ahora se sirven desde el RAG v3 remoto en el Epyc/Sonora, vía RAG_V3_URL.
# _inicializar_chroma() se deja definida por si se necesita reactivar.
print("--- MIA HERRAMIENTAS: usando RAG v3 remoto (Sonora) ---")
print(f"    Endpoint: {RAG_V3_URL}")
try:
    _test = requests.get(RAG_V3_URL.replace('/buscar', '/status'), timeout=5)
    if _test.ok:
        _status = _test.json()
        print(f"    ✅ RAG remoto activo: {_status.get('chunks_totales', '?')} chunks indexados")
    else:
        print(f"    ⚠️ RAG remoto respondió con estado {_test.status_code}")
except Exception as _e:
    print(f"    ❌ RAG remoto no disponible: {_e}")


# ============================================================================
# BÚSQUEDA PBIP — migrada al RAG v3 remoto (Epyc/Sonora), agosto 2026
# ============================================================================

def buscar_pbip_hibrido(query: str, k: int = 5, parte: str = None,
                        estrategia: str = 'auto') -> list[dict]:
    """
    Consulta el RAG v3 remoto. Se conserva el nombre y firma originales
    (incluyendo 'parte' y 'estrategia', ya sin uso real) para no romper
    a mia_core.py y test_hibrido.py, que importan esta función directamente.
    Retorna list[{'text','metadata'}], mismo contrato que la versión local.
    """
    try:
        resp = requests.post(
            RAG_V3_URL,
            json={"pregunta": query, "k": k},
            headers={"Authorization": f"Bearer {RAG_V3_TOKEN}"},
            timeout=30
        )
        resp.raise_for_status()
        data = resp.json()
        resultados_raw = data.get("resultados", [])

        resultados = []
        for r in resultados_raw:
            resultados.append({
                "text": r.get("texto", ""),
                "metadata": {
                    "parte": "",
                    "seccion": r.get("seccion") or r.get("documento", "PBIP"),
                    "documento": r.get("documento", ""),
                },
                "final_score": r.get("score_rerank", 0.0),
            })
        return resultados

    except Exception as e:
        print(f"Error buscar_pbip_hibrido (RAG remoto): {e}")
        return []


# ============================================================================
# LLM (Ollama)
# ============================================================================

def consultar_ollama(prompt: str, temperature: float = 0.2, num_ctx: int = 16384) -> str:
    url = "http://localhost:11434/api/generate"
    payload = {
        "model": "qwen2.5:7b-instruct-q4_K_M",
        "prompt": prompt,
        "stream": False,
        "options": {"num_ctx": num_ctx, "temperature": temperature, "num_gpu": 99}
    }
    try:
        r = requests.post(url, json=payload, timeout=600)
        return r.json().get('response', 'Sin respuesta')
    except requests.exceptions.Timeout:
        return "⏱️ Error IA: El modelo tardó demasiado en responder. Intenta con una pregunta más corta o verifica que Ollama no esté sobrecargado."
    except Exception as e:
        return f"Error IA: {e}"


# ============================================================================
# EXTRACCIÓN DE TEXTO UNIVERSAL
# ============================================================================

def extraer_texto_universal(ruta_archivo: str) -> str:
    if not os.path.exists(ruta_archivo):
        return "Error: El archivo no existe."

    extension = os.path.splitext(ruta_archivo)[1].lower()
    texto_extraido = ""

    try:
        if extension == '.docx' and DOCX_AVAILABLE:
            print(f"📝 Word: {os.path.basename(ruta_archivo)}")
            doc = DocxDocument(ruta_archivo)
            parrafos = [p.text for p in doc.paragraphs if p.text]
            texto_extraido = "\n".join(parrafos)
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        if celda.text.strip():
                            texto_extraido += f"\n{celda.text.strip()}"

        elif extension in ['.jpg', '.jpeg', '.png']:
            print(f"📷 Imagen: {os.path.basename(ruta_archivo)}")
            if OCR_AVAILABLE:
                imagen = Image.open(ruta_archivo)
                texto_extraido = pytesseract.image_to_string(imagen, lang='spa')
            else:
                return "Error: OCR no disponible."

        elif extension == '.pdf':
            print(f"📄 PDF: {os.path.basename(ruta_archivo)}")
            with pdfplumber.open(ruta_archivo) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        texto_extraido += page_text + "\n"

            if len(texto_extraido.strip()) < 200 and OCR_AVAILABLE:
                print("📷 PDF plano, aplicando OCR...")
                paginas = convert_from_path(ruta_archivo)
                texto_ocr = ""
                for img in paginas:
                    texto_ocr += pytesseract.image_to_string(img, lang='spa') + "\n"
                texto_extraido = texto_ocr

        else:
            return f"Error: Formato {extension} no soportado."

        if len(texto_extraido.strip()) > 10:
            return texto_extraido
        else:
            return "Error: No se extrajo texto suficiente."

    except Exception as e:
        return f"Error interno: {str(e)}"


# ============================================================================
# HERRAMIENTA CONSULTAR PBIP (VERSIÓN HÍBRIDA)
# ============================================================================

def buscar_analisis_situacion(descripcion: str) -> dict:
    """
    Consulta la memoria de análisis de situaciones (AnalisisSituacion en
    MIA 2.0/Epyc) antes de generar una respuesta nueva. Si hay una situación
    similar ya validada, se reutiliza directamente — evita regenerar y
    elimina el riesgo de variabilidad/alucinación del LLM en preguntas
    ya resueltas antes.
    Retorna {'encontrado': bool, 'analisis': str, 'similitud': float, ...}
    """
    try:
        resp = requests.post(
            f"{ANALISIS_MEMORIA_URL}/buscar",
            json={"descripcion": descripcion},
            timeout=10
        )
        resp.raise_for_status()
        return resp.json()
    except Exception as e:
        print(f"Error consultando memoria de análisis: {e}")
        return {"encontrado": False}


def guardar_analisis_situacion(descripcion: str, analisis: str, tipo_analisis: str = "pbip_consulta",
                                 normas_citadas: list = None) -> bool:
    """Guarda un análisis validado en la memoria compartida (AnalisisSituacion)."""
    try:
        resp = requests.post(
            f"{ANALISIS_MEMORIA_URL}/guardar",
            json={
                "descripcion": descripcion,
                "analisis": analisis,
                "tipo_analisis": tipo_analisis,
                "normas_citadas": normas_citadas,
                "origen": "compass",
            },
            timeout=10
        )
        resp.raise_for_status()
        return resp.json().get("guardado", False)
    except Exception as e:
        print(f"Error guardando en memoria de análisis: {e}")
        return False


def herramienta_consultar_pbip(tema: str, k: int = 5, parte: str = None) -> str:
    """
    Consulta el RAG marítimo v3 (Epyc/Sonora) — reranker multilingüe,
    chunker corregido, corpus ampliado (PBIP, ISM, MARPOL, STCW, IMDG,
    legislación mexicana). Reemplaza el motor local viejo (buscar_pbip_hibrido)
    a partir de agosto 2026.
    Esta función es usada por:
      - mia_core.py (consultas del auditor por WhatsApp)
      - mia_documentos.py (análisis de documentos subidos por navieras)
    """
    try:
        resp = requests.post(
            RAG_V3_URL,
            json={"pregunta": tema, "k": k},
            headers={"Authorization": f"Bearer {RAG_V3_TOKEN}"},
            timeout=30
        )
        resp.raise_for_status()
        data = resp.json()
        resultados = data.get("resultados", [])

        if not resultados:
            return f"No encontré información sobre '{tema}' en el código PBIP."

        lineas = []
        for r in resultados:
            documento = r.get('documento', 'PBIP')
            seccion = r.get('seccion') or ''
            score = r.get('score_rerank', 0)

            ref = f"{documento}"
            if seccion:
                ref += f" | {seccion}"
            ref += f" (relevancia: {score:.2f})"

            contenido = r['texto'][:1200] if len(r['texto']) > 600 else r['texto']
            lineas.append(f"[{ref}]\n{contenido}")

        return "\n\n---\n\n".join(lineas)

    except requests.exceptions.RequestException as e:
        return f"Error consultando PBIP (RAG remoto no disponible): {e}"
    except Exception as e:
        return f"Error consultando PBIP: {e}"


# ============================================================================
# BÚSQUEDA UNIFICADA — PBIP + LEGISLACIÓN + FUTUROS DOMINIOS
# ============================================================================

def buscar_legislacion(query: str, k: int = 6) -> list:
    """
    Migrado al RAG v3 remoto (Epyc/Sonora) — agosto 2026.
    Mantiene el mismo contrato de retorno (list[{'text','metadata'}]) que
    espera mia_core.py, para no romper _modo_consulta_legislacion.
    """
    try:
        resp = requests.post(
            RAG_V3_URL,
            json={"pregunta": query, "k": k},
            headers={"Authorization": f"Bearer {RAG_V3_TOKEN}"},
            timeout=30
        )
        resp.raise_for_status()
        data = resp.json()
        resultados_raw = data.get("resultados", [])

        resultados = []
        for r in resultados_raw:
            resultados.append({
                "text": r.get("texto", ""),
                "metadata": {
                    "fuente": r.get("documento", "Legislación"),
                    "articulo": r.get("seccion") or "",
                }
            })
        return resultados

    except Exception as e:
        print(f"Error busqueda legislacion (RAG remoto): {e}")
        return []

# (buscar_conocimiento y herramienta_consultar_conocimiento eliminadas — sin uso, ago 2026)


# ============================================================================
# DJANGO ORM - ESTADO DE EXPEDIENTES
# ============================================================================

def herramienta_consultar_estado(naviera_nombre: str = None, buque_nombre: str = None, omi: str = None) -> str:
    try:
        if omi:
            try:
                b = Buque.objects.get(OMI=omi)
                return _formato_buque(b)
            except Buque.DoesNotExist:
                return f"❌ No encontré buque con OMI: {omi}"

        if buque_nombre:
            buques = Buque.objects.filter(nombre_buque__icontains=buque_nombre)
            if buques.count() == 1:
                return _formato_buque(buques.first())
            elif buques.count() > 1:
                return f"⚠️ {buques.count()} buques coinciden:\n" + "\n".join([f"• {b.nombre_buque} (OMI:{b.OMI})" for b in buques[:5]])

        if naviera_nombre:
            navieras = Naviera.objects.filter(nombre_empresa__icontains=naviera_nombre)
            if navieras.count() == 1:
                return _formato_naviera(navieras.first())
            elif navieras.count() > 1:
                return f"⚠️ {navieras.count()} navieras:\n" + "\n".join([f"• {n.nombre_empresa}" for n in navieras[:5]])

        return "Especifica naviera, buque u OMI."

    except Exception as e:
        return f"Error consultando estado: {e}"


def _formato_buque(buque) -> str:
    total_pbip = PuntoPBIP.objects.count()
    pbip_subidos = RequisitoBuque.objects.filter(buque=buque, categoria='DOCUMENTAL').count()
    pct = int((pbip_subidos / total_pbip) * 100) if total_pbip else 0

    return f"""🚢 *{buque.nombre_buque}*
📋 OMI: {buque.OMI}
🏢 {buque.naviera.nombre_empresa}
📊 PBIP: {pbip_subidos}/{total_pbip} ({pct}%)"""


def _formato_naviera(naviera) -> str:
    buques = Buque.objects.filter(naviera=naviera)
    lineas = [f"🏢 *{naviera.nombre_empresa}*\n🚢 Buques: {buques.count()}"]

    total_pbip = PuntoPBIP.objects.count()
    for b in buques:
        pbip_subidos = RequisitoBuque.objects.filter(buque=b, categoria='DOCUMENTAL').count()
        pct = int((pbip_subidos / total_pbip) * 100) if total_pbip else 0
        lineas.append(f"• {b.nombre_buque} (OMI:{b.OMI}): {pct}%")

    return "\n".join(lineas)


def herramienta_reporte_global() -> str:
    buques = Buque.objects.all()
    total_pbip = PuntoPBIP.objects.count()
    total_admin = 6

    lineas = ["🧭 *COMPASS - ESTADO GLOBAL*\n"]

    for buque in buques:
        pbip_subidos = RequisitoBuque.objects.filter(
            buque=buque, 
            categoria='DOCUMENTAL'
        ).count()
        pct_pbip = int((pbip_subidos / total_pbip) * 100) if total_pbip else 0

        naviera = buque.naviera
        admin_subidos = RequisitoBuque.objects.filter(
            naviera=naviera,
            buque__isnull=True,
            categoria='ADMINISTRATIVO'
        ).count()
        pct_admin = int((admin_subidos / total_admin) * 100)

        estado_pbip = "✅" if pct_pbip == 100 else f"{pct_pbip}%"
        estado_admin = "✅" if pct_admin == 100 else f"{pct_admin}%"

        lineas.append(
            f"• {buque.nombre_buque[:20]} (OMI:{buque.OMI})\n"
            f"  📊 PBIP: {estado_pbip} | 🏢 Admin: {estado_admin}"
        )

    lineas.append("\n🏢 *RESUMEN POR NAVIERA:*")
    navieras = Naviera.objects.all()
    for nav in navieras:
        admin_count = RequisitoBuque.objects.filter(
            naviera=nav,
            buque__isnull=True,
            categoria='ADMINISTRATIVO'
        ).count()
        buques_nav = Buque.objects.filter(naviera=nav)
        pbip_total = sum(
            RequisitoBuque.objects.filter(buque=b, categoria='DOCUMENTAL').count()
            for b in buques_nav
        )
        lineas.append(
            f"• {nav.nombre_empresa[:25]}: "
            f"Admin {admin_count}/6 | PBIP {pbip_total}/{total_pbip * buques_nav.count()}"
        )

    lineas.append(f"\n_Total: {buques.count()} buques | {navieras.count()} navieras_")
    return "\n".join(lineas)


# ============================================================================
# WHATSAPP
# ============================================================================

def enviar_whatsapp_jid(jid: str, mensaje: str) -> bool:
    if '@' not in jid:
        jid = f"{jid}@s.whatsapp.net"

    try:
        requests.post(
            "http://100.112.139.108:9000/enviar",
            json={"jid": jid, "mensaje": mensaje},
            timeout=10
        )
    except:
        pass

    return True


def enviar_whatsapp_numero(numero: str, mensaje: str) -> bool:
    try:
        requests.post(
            "http://100.112.139.108:9000/enviar",
            json={"numero": numero, "mensaje": mensaje},
            timeout=10
        )
        return True
    except Exception as e:
        print(f"❌ Error WhatsApp: {e}")
        return False

# Agregar a mia_herramientas.py

def enviar_opr_notificacion(jid: str, mensaje: str) -> bool:
    """Envía notificación por OPR Gateway (segundo Baileys)."""
    if '@' not in jid:
        jid = f"{jid}@s.whatsapp.net"
    
    try:
        r = requests.post(
            "http://localhost:9001/enviar",
            json={"jid": jid, "mensaje": mensaje},
            timeout=15
        )
        return r.status_code == 200
    except Exception as e:
        print(f"❌ Error OPR Gateway: {e}")
        return False