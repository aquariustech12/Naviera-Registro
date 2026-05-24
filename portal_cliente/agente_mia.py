import os
import re
import json
import requests
import pdfplumber
from datetime import datetime
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings
from PIL import Image
from docx import Document

# OCR
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

CHROMA_PATH = "/home/julian/Naviera-Registro/scripts/chroma_db"

print(f"--- MIA INICIANDO ---")
print(f"OCR: {'✅' if OCR_AVAILABLE else '❌'}")

embeddings = OllamaEmbeddings(model="nomic-embed-text")

try:
    vector_db = Chroma(
        persist_directory=CHROMA_PATH,
        embedding_function=embeddings,
        collection_name="auditoria_pbip"
    )
    count = vector_db._collection.count()
    print(f"✅ Base de conocimiento (PBIP): {count} documentos")
except Exception as e:
    print(f"❌ ERROR cargando Chroma: {e}")
    vector_db = None

def extraer_texto_universal(ruta_archivo):
    """
    Detecta la extensión del archivo y extrae el texto ya sea un PDF, 
    una imagen de WhatsApp (JPEG/PNG) o un documento de Word (DOCX).
    """
    if not os.path.exists(ruta_archivo):
        return "Error: El archivo no existe en la ruta especificada."

    extension = os.path.splitext(ruta_archivo)[1].lower()
    texto_extraido = ""

    try:
        # CASO 1: EL ARCHIVO ES UN WORD (.DOCX)
        if extension == '.docx':
            print(f"📝 Procesando documento de Word: {os.path.basename(ruta_archivo)}")
            doc = Document(ruta_archivo)
            parrafos = [p.text for p in doc.paragraphs if p.text]
            texto_extraido = "\n".join(parrafos)
            
            # También extraemos texto dentro de tablas de Word si existen
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        if celda.text.strip():
                            texto_extraido += f"\n{celda.text.strip()}"

        # CASO 2: EL ARCHIVO ES UNA FOTO / IMAGEN (.JPG, .JPEG, .PNG)
        elif extension in ['.jpg', '.jpeg', '.png']:
            print(f"📷 Procesando imagen directa (FOTO): {os.path.basename(ruta_archivo)}")
            if OCR_AVAILABLE:
                # Usamos Pillow para abrir la imagen y Tesseract directo a la vena
                imagen = Image.open(ruta_archivo)
                texto_extraido = pytesseract.image_to_string(imagen, lang='spa')
            else:
                return "Error: Se subió una imagen pero Tesseract no está disponible en el servidor."

        # CASO 3: EL FORMATO UNIVERSAL (PDF)
        elif extension == '.pdf':
            print(f"📄 Procesando PDF: {os.path.basename(ruta_archivo)}")
            with pdfplumber.open(ruta_archivo) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        texto_extraido += page_text + "\n"
            
            # Si el PDF era solo una imagen escaneada (menos de 200 letras), metemos el OCR que ya tenías
            if len(texto_extraido.strip()) < 200 and OCR_AVAILABLE:
                print("📷 PDF plano detectado. Aplicando OCR por páginas...")
                paginas_imagenes = convert_from_path(ruta_archivo)
                texto_ocr = ""
                for img in paginas_imagenes:
                    texto_ocr += pytesseract.image_to_string(img, lang='spa') + "\n"
                texto_extraido = texto_ocr

        # CASO 4: FORMATO NO SOPORTADO
        else:
            return f"Error: El formato {extension} no está soportado actualmente por MIA."

        # Validación final del texto obtenido
        if len(texto_extraido.strip()) > 10:
            return texto_extraido
        else:
            return "Error: El archivo se leyó pero no se pudo extraer texto suficiente."

    except Exception as e:
        print(f"❌ Error crítico procesando {extension}: {e}")
        return f"Error interno del servidor al procesar el archivo: {str(e)}"

def consultar_ollama(prompt, temperature=0.2):
    url = "http://localhost:11434/api/generate"
    payload = {
        "model": "qwen2.5:14b",
        "prompt": prompt,
        "stream": False,
        "options": {"num_ctx": 16384, "temperature": temperature}  # Aumentado contexto para análisis completo
    }
    try:
        r = requests.post(url, json=payload, timeout=300)
        return r.json().get('response', 'Sin respuesta')
    except Exception as e:
        return f"Error IA: {e}"

def consultar_pbip_relevante(texto_documento):
    if vector_db is None:
        return "Base PBIP no disponible"
    try:
        query = texto_documento[:4000]  # Más contexto para mejor búsqueda
        docs = vector_db.similarity_search(query, k=8)  # Más resultados
        if not docs:
            return "No se encontraron artículos relevantes"
        contextos = []
        for doc in docs:
            meta = doc.metadata
            ref = f"{meta.get('seccion', 'PBIP')}"
            if meta.get('articulo'):
                ref += f", {meta['articulo']}"
            content = doc.page_content[:1000]  # Más contenido por artículo
            contextos.append(f"[{ref}]\n{content}")
        return "\n---\n".join(contextos)
    except Exception as e:
        return f"Error consultando PBIP: {e}"

def ejecutar_analisis_mia(documento_obj, jid_remitente=None):
    try:
        print(f"\n📄 Analizando: {documento_obj.nombre_documento}")
        texto = extraer_texto_universal(documento_obj.archivo.path)
        
        if texto.startswith("ERROR"):
            return enviar_whatsapp_mia(f"🤖 MIA - ERROR\n📄 {documento_obj.nombre_documento}\n❌ {texto}", jid_remitente)

        # Filtro de documentos administrativos (No bloquean, no auditan)
        admin_keywords = ['factura', 'cotizacion', 'pago', 'rfc', 'domicilio', 'fgmp-fc-01']
        es_admin = any(key in documento_obj.nombre_documento.lower() for key in admin_keywords)

        if es_admin:
            # Respuesta inmediata para administrativos con tu estructura
            mensaje_admin = f"""🤖 *MIA - REGISTRO*
📄 *Documento:* {documento_obj.nombre_documento}
🏷️ *Tipo identificado:* Administrativo
🔍 *Análisis PBIP:* No requiere evaluación técnica. Documento registrado para control administrativo.
✅/❌ *Dictamen:* NO APLICA PBIP"""
            return enviar_whatsapp_mia(mensaje_admin, jid_remitente)

        # --- Análisis PBIP Técnico (Tu lógica original intacta) ---
        fecha_hoy = datetime.now().strftime("%d/%m/%Y")
        contexto_pbip = consultar_pbip_relevante(texto)
        
        prompt = f"""Eres MIA, auditor experto en el Código PBIP.
FECHA DE HOY: {fecha_hoy}
DOCUMENTO RECIBIDO: "{documento_obj.nombre_documento}"
CONTENIDO: {texto[:5000]}
ARTÍCULOS RELEVANTES: {contexto_pbip}

INSTRUCCIONES:
1. Identifica el tipo de documento.
2. Si es PBIP, extrae Titular, Folio, Expedición y Vigencia.
3. Analiza contra los artículos de Chroma.
4. Si no es PBIP, indica "No aplica".

FORMATO DE RESPUESTA (ESTRICTO):
🤖 *MIA - AUDITORÍA*
📄 *Documento:* {documento_obj.nombre_documento}
🏷️ *Tipo identificado:* [OPB/PFSO/Plan de protección/Administrativo/etc.]
👤 *Titular/Responsable:* [nombre o No encontrado]
📜 *Folio/Certificado:* [número o No encontrado]
📅 *Expedición:* [fecha o No encontrado]
⏰ *Vigencia:* [fecha o No aplica]
🔍 *Análisis PBIP:* [Evaluación detallada]
📜 *Base Legal:* [Artículos aplicables]
✅/❌ *Dictamen:* [CUMPLE / NO CUMPLE / NO APLICA PBIP]
💡 *Recomendación:* [solo si hay deficiencias]"""

        resultado = consultar_ollama(prompt, temperature=0.2)
        enviar_whatsapp_mia(resultado, jid_remitente)
        return True

    except Exception as e:
        enviar_whatsapp_mia(f"🤖 MIA - ERROR\n📄 {documento_obj.nombre_documento}\n❌ {str(e)}", jid_remitente)
        return False

def enviar_whatsapp_mia(mensaje, jid=None):
    try:
        destino = jid if jid else "5216444475422@s.whatsapp.net"
        if '@' not in destino:
            destino = f"{destino}@s.whatsapp.net"
            
        print(f"📤 Enviando WhatsApp a: {destino}")
            
        requests.post("http://localhost:9000/enviar", 
                     json={"jid": destino, "mensaje": mensaje}, 
                     timeout=10)
        return True
    except Exception as e:
        print(f"❌ Error enviando WhatsApp: {e}")
        return False

def validar_correspondencia(nombre_esperado, texto_documento):
    """
    Valida si el documento corresponde al tipo esperado, con tolerancia a errores de OCR
    y variaciones en el nombre.
    """
    # Si el texto es demasiado corto, asumimos que es ilegible -> no corresponde
    if len(texto_documento.strip()) < 50:
        return False, "El documento no contiene texto suficiente (puede estar escaneado sin OCR o ser ilegible)"

    prompt = f"""Eres un asistente experto en documentación marítima. Determina si el documento corresponde al tipo esperado.

TIPO ESPERADO: "{nombre_esperado}"

CONTENIDO DEL DOCUMENTO:
{texto_documento[:2000]}

Instrucciones:
- Responde ÚNICAMENTE con un objeto JSON válido.
- Sé permisivo: si el documento trata sobre el mismo tema (aunque el nombre no sea idéntico o falten detalles), responde true.
- Solo responde false si el documento es claramente de otro tema (ej. factura, identificación personal, contrato de otro rubro).
- La razón debe ser breve y en español.

Ejemplo:
{{"corresponde": true, "razon": "El documento contiene información sobre competencia de oficial de protección, aunque el título varía"}}

Ahora evalúa:
{{"corresponde": true/false, "razon": "..."}}"""

    try:
        respuesta = consultar_ollama(prompt, temperature=0.0)
        # Limpiar posibles caracteres no JSON
        respuesta = respuesta.strip()
        if respuesta.startswith('```json'):
            respuesta = respuesta[7:]
        if respuesta.endswith('```'):
            respuesta = respuesta[:-3]
        data = json.loads(respuesta)
        return data.get("corresponde", False), data.get("razon", "No se pudo determinar")
    except Exception as e:
        print(f"Error en validación: {e}")
        # Si falla el parsing, asumimos que no corresponde
        return False, f"Error interno al validar: {str(e)}"